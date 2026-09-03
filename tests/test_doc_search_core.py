# -*- coding: utf-8 -*-
"""ტესტები — დოკუმენტებში ძებნა: ამოღება, ინდექსი, ძებნა, ამონარიდი.

ყველა ფაილი ტესტში იქმნება (tmp_path) — რეპოში სატესტო დოკუმენტი არ დევს.
"""

import os

import pytest

import tools.doc_search_core as ds


def _tr(key, **fmt):
    return key


def _write(tmp_path, name, text):
    p = tmp_path / name
    p.write_text(text, encoding="utf-8")
    return str(p)


# ---- parse_query -----------------------------------------------------------
def test_phrase_allows_line_break_between_words():
    """PDF-ში ორ ხაზად გატეხილი ფრაზა მაინც უნდა მოიძებნოს."""
    rx = ds.parse_query("ნაკვეთის საზღვარი", ds.MODE_PHRASE)[0]
    assert rx.search("…მიწის ნაკვეთის\n   საზღვარი გადის…")


def test_phrase_is_substring_not_word_bounded():
    """ქართული აგლუტინაციურია — „ნაკვეთ“ უნდა იპოვოს „ნაკვეთებში“."""
    rx = ds.parse_query("ნაკვეთ", ds.MODE_PHRASE)[0]
    assert rx.search("ორივე ნაკვეთებში")


def test_any_mode_splits_on_pipe_and_space():
    rxs = ds.parse_query("საზღვარი | მიჯნა", ds.MODE_ANY)
    assert len(rxs) == 2


def test_empty_query_gives_no_regexes():
    assert ds.parse_query("   ", ds.MODE_PHRASE) == []


def test_special_characters_are_escaped():
    """„(“ regex-ს არ უნდა გატეხოს — ტექსტად უნდა ჩაითვალოს."""
    rx = ds.parse_query("მუხლი 3(ა)", ds.MODE_PHRASE)[0]
    assert rx.search("იხ. მუხლი 3(ა) —")


# ---- anchor_for ------------------------------------------------------------
def test_anchor_for_finds_the_page():
    anchors = [(0, 1), (100, 2), (250, 3)]
    assert ds.anchor_for(anchors, 0) == 1
    assert ds.anchor_for(anchors, 99) == 1
    assert ds.anchor_for(anchors, 100) == 2
    assert ds.anchor_for(anchors, 300) == 3


def test_anchor_for_empty():
    assert ds.anchor_for([], 5) is None


# ---- make_snippet ----------------------------------------------------------
def test_snippet_highlight_points_at_the_match():
    text = "a" * 300 + " ნაკვეთის საზღვარი " + "b" * 300
    start = text.index("ნაკვეთის")
    end = start + len("ნაკვეთის საზღვარი")
    snippet, hl = ds.make_snippet(text, [(start, end)])
    assert len(hl) == 1
    s, e = hl[0]
    assert snippet[s:e] == "ნაკვეთის საზღვარი"


def test_snippet_collapses_whitespace_and_still_highlights():
    text = "წინა ტექსტი ნაკვეთის\n\n  საზღვარი შემდეგი"
    start = text.index("ნაკვეთის")
    end = text.index("საზღვარი") + len("საზღვარი")
    snippet, hl = ds.make_snippet(text, [(start, end)])
    assert "\n" not in snippet
    s, e = hl[0]
    assert snippet[s:e] == "ნაკვეთის საზღვარი"


def test_snippet_marks_truncated_edges():
    text = "x" * 500 + " სიტყვა " + "y" * 500
    start = text.index("სიტყვა")
    snippet, _ = ds.make_snippet(text, [(start, start + len("სიტყვა"))])
    assert snippet.startswith("…") and snippet.endswith("…")


def test_two_matches_in_one_paragraph_give_one_snippet_two_marks():
    """რეგრესია: ერთ პარაგრაფში ორი დამთხვევა ორ იდენტურ ამონარიდს ბადებდა და
    მონიშვნა ორივეჯერ პირველ დამთხვევაზე იდგა."""
    text = ("ნაკვეთის საზღვარი დადგენილია ბუნებრივი ნიშნულებით. "
            "სამხრეთიდან ნაკვეთის საზღვარი ემიჯნება ტყის ფონდს.")
    rx = ds.parse_query("ნაკვეთის საზღვარი", ds.MODE_PHRASE)[0]
    spans = [(m.start(), m.end()) for m in rx.finditer(text)]
    assert len(spans) == 2

    groups = ds.group_spans(spans)
    assert len(groups) == 1                      # ერთ ამონარიდში უნდა გაერთიანდეს

    snippet, hl = ds.make_snippet(text, groups[0])
    assert len(hl) == 2                          # ორივე მონიშნულია
    assert [snippet[a:b] for a, b in hl] == ["ნაკვეთის საზღვარი"] * 2
    assert hl[0][0] != hl[1][0]                  # და სხვადასხვა ადგილას დგას


def test_far_apart_matches_stay_separate():
    assert len(ds.group_spans([(0, 8), (1008, 1016)])) == 2


def test_dense_repetition_does_not_make_one_endless_snippet():
    text = "საზღვარი " * 200
    rx = ds.parse_query("საზღვარი", ds.MODE_PHRASE)[0]
    groups = ds.group_spans([(m.start(), m.end()) for m in rx.finditer(text)])
    assert len(groups) > 1
    for g in groups:
        assert g[-1][1] - g[0][0] <= ds.MAX_SNIPPET_SPAN


# ---- ინდექსი + ძებნა -------------------------------------------------------
def test_index_and_search_roundtrip(tmp_path):
    _write(tmp_path, "a.txt", "მიწის ნაკვეთის საზღვარი გადის სამხრეთით.")
    _write(tmp_path, "b.txt", "აქ სხვა შინაარსია, საზღვარი არსად წერია.")
    _write(tmp_path, "c.txt", "ნაკვეთის საზღვარი ორჯერ: ნაკვეთის საზღვარი.")
    db = str(tmp_path / "idx.db")

    done = {}
    ds.index_folder(str(tmp_path), False, db, lambda m: None,
                    lambda r: done.update(r), _tr)
    assert done["total"] == 3
    assert done["new"] == 3

    res = ds.search(db, "ნაკვეთის საზღვარი", ds.MODE_PHRASE)
    names = [r["name"] for r in res]
    assert names == ["c.txt", "a.txt"]          # მეტ დამთხვევიანი წინ
    assert res[0]["count"] == 2
    assert res[1]["count"] == 1
    assert "b.txt" not in names


def test_second_index_uses_the_cache(tmp_path):
    _write(tmp_path, "a.txt", "ნაკვეთის საზღვარი")
    db = str(tmp_path / "idx.db")
    for _ in range(2):
        done = {}
        ds.index_folder(str(tmp_path), False, db, lambda m: None,
                        lambda r: done.update(r), _tr)
    assert done["cached"] == 1                  # მეორე ჯერზე თავიდან არ წაუკითხავს
    assert done["new"] == 0


def test_changed_file_is_reindexed(tmp_path):
    p = _write(tmp_path, "a.txt", "ძველი შინაარსი")
    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    assert ds.search(db, "ახალი", ds.MODE_PHRASE) == []

    os.utime(p, (0, 0))                         # mtime იცვლება
    with open(p, "w", encoding="utf-8") as fh:
        fh.write("ახალი შინაარსი აქაა")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    assert len(ds.search(db, "ახალი", ds.MODE_PHRASE)) == 1


def test_deleted_file_drops_out_of_the_index(tmp_path):
    p = _write(tmp_path, "a.txt", "ნაკვეთის საზღვარი")
    _write(tmp_path, "b.txt", "ნაკვეთის საზღვარი")
    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    assert len(ds.search(db, "საზღვარი", ds.MODE_PHRASE)) == 2

    os.remove(p)
    done = {}
    ds.index_folder(str(tmp_path), False, db, lambda m: None,
                    lambda r: done.update(r), _tr)
    assert done["removed"] == 1
    assert len(ds.search(db, "საზღვარი", ds.MODE_PHRASE)) == 1


def test_all_mode_needs_every_word(tmp_path):
    _write(tmp_path, "both.txt", "ნაკვეთი აქ, ხოლო საზღვარი იქ.")
    _write(tmp_path, "one.txt", "მხოლოდ ნაკვეთი.")
    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)

    res_all = ds.search(db, "ნაკვეთი საზღვარი", ds.MODE_ALL)
    assert [r["name"] for r in res_all] == ["both.txt"]

    res_any = ds.search(db, "ნაკვეთი საზღვარი", ds.MODE_ANY)
    assert sorted(r["name"] for r in res_any) == ["both.txt", "one.txt"]


def test_recursive_flag(tmp_path):
    sub = tmp_path / "sub"
    sub.mkdir()
    _write(tmp_path, "top.txt", "საზღვარი")
    (sub / "deep.txt").write_text("საზღვარი", encoding="utf-8")

    assert len(ds.iter_documents(str(tmp_path), recursive=False)) == 1
    assert len(ds.iter_documents(str(tmp_path), recursive=True)) == 2


def test_word_temp_files_are_ignored(tmp_path):
    _write(tmp_path, "real.docx", "x")
    _write(tmp_path, "~$real.docx", "x")
    found = [os.path.basename(p) for p in ds.iter_documents(str(tmp_path))]
    assert found == ["real.docx"]


def test_index_reports_progress_and_cancel(tmp_path):
    for i in range(5):
        _write(tmp_path, "f%d.txt" % i, "ნაკვეთი %d" % i)
    db = str(tmp_path / "idx.db")

    prog = []
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None,
                    _tr, progress=lambda i, n: prog.append((i, n)))
    assert prog[-1] == (5, 5)

    done = {}
    ds.index_folder(str(tmp_path), False, str(tmp_path / "idx2.db"),
                    lambda m: None, lambda r: done.update(r), _tr,
                    cancel=lambda: True, force=True)
    assert done.get("cancelled") is True


def test_empty_folder_and_missing_folder_report_errors(tmp_path):
    empty = tmp_path / "empty"
    empty.mkdir()
    done = {}
    ds.index_folder(str(empty), False, str(tmp_path / "i.db"), lambda m: None,
                    lambda r: done.update(r), _tr)
    assert done["error"] == "no_docs"

    done = {}
    ds.index_folder(str(tmp_path / "nope"), False, str(tmp_path / "i.db"),
                    lambda m: None, lambda r: done.update(r), _tr)
    assert done["error"] == "warn_folder"


def test_scanned_like_file_is_flagged_not_searchable(tmp_path):
    """ტექსტის გარეშე ფაილი „empty“-ად უნდა მოინიშნოს და ძებნაში არ მოხვდეს."""
    _write(tmp_path, "scan.txt", "  ")
    _write(tmp_path, "ok.txt", "ნაკვეთის საზღვარი გადის სამხრეთით.")
    db = str(tmp_path / "idx.db")
    done = {}
    ds.index_folder(str(tmp_path), False, db, lambda m: None,
                    lambda r: done.update(r), _tr)
    assert done["empty"] == 1

    stats = ds.index_stats(db)
    assert stats["ok"] == 1 and stats["empty"] == 1
    assert [os.path.basename(p) for p in stats["empty_files"]] == ["scan.txt"]


def test_search_hits_are_capped_but_count_is_full(tmp_path):
    # დამთხვევებს ვაშორებთ, რომ თითო თავის ამონარიდში მოხვდეს
    _write(tmp_path, "many.txt", ("საზღვარი" + "-" * 600) * 40)
    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    res = ds.search(db, "საზღვარი", ds.MODE_PHRASE, max_hits_per_file=10)
    assert res[0]["count"] == 40                # სრული რაოდენობა უცვლელია
    assert len(res[0]["hits"]) == 10            # ნაჩვენები კი შეზღუდული
    assert res[0]["truncated"] is True


def test_search_labels_carry_the_location(tmp_path):
    _write(tmp_path, "a.txt", "ნაკვეთის საზღვარი")
    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    hit = ds.search(db, "საზღვარი", ds.MODE_PHRASE)[0]["hits"][0]
    assert hit["label"] == 1


def test_unsupported_extension_is_skipped(tmp_path):
    res = ds.extract_text(str(_write(tmp_path, "x.xyz", "abc")))
    assert res["status"] == "skipped"


# ---- .docx (python-docx-ის არსებობისას) ------------------------------------
@pytest.mark.skipif(ds.docx is None, reason="python-docx is not installed")
def test_docx_table_cells_are_indexed(tmp_path):
    """კადასტრულ დოკუმენტებში ტექსტის დიდი ნაწილი ცხრილშია — უნდა მოიძებნოს."""
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("სათაური")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "მაჩვენებელი"
    table.cell(0, 1).text = "მნიშვნელობა"
    table.cell(1, 0).text = "ნაკვეთის საზღვარი"
    table.cell(1, 1).text = "დაზუსტებულია"
    path = str(tmp_path / "t.docx")
    doc.save(path)

    res = ds.extract_text(path)
    assert res["status"] == "ok"
    assert "ნაკვეთის საზღვარი" in res["text"]
    assert "დაზუსტებულია" in res["text"]
    assert "სათაური" in res["text"]


@pytest.mark.skipif(ds.docx is None, reason="python-docx is not installed")
def test_docx_keeps_document_order(tmp_path):
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("პირველი")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "მეორე"
    doc.add_paragraph("მესამე")
    path = str(tmp_path / "o.docx")
    doc.save(path)

    text = ds.extract_text(path)["text"]
    assert text.index("პირველი") < text.index("მეორე") < text.index("მესამე")


# ---- .pdf (PyMuPDF-ის არსებობისას) -----------------------------------------
@pytest.mark.skipif(ds.fitz is None, reason="PyMuPDF is not installed")
def test_pdf_without_a_text_layer_is_flagged_empty(tmp_path):
    """ცარიელგვერდიანი PDF = დასკანერებულის იმიტაცია — „empty“-ად უნდა მოინიშნოს."""
    doc = ds.fitz.open()
    for _ in range(3):
        doc.new_page()
    path = str(tmp_path / "scan.pdf")
    doc.save(path)
    doc.close()
    assert ds.extract_text(path)["status"] == "empty"


@pytest.mark.skipif(ds.fitz is None, reason="PyMuPDF is not installed")
def test_short_pdf_with_real_text_is_not_flagged(tmp_path):
    """მოკლე, მაგრამ ნამდვილი ტექსტი „empty“ არ არის."""
    doc = ds.fitz.open()
    doc.new_page().insert_text((72, 72), "sazghvari gadis", fontsize=12)
    path = str(tmp_path / "short.pdf")
    doc.save(path)
    doc.close()
    assert ds.extract_text(path)["status"] == "ok"


@pytest.mark.skipif(ds.fitz is None, reason="PyMuPDF is not installed")
def test_pdf_page_numbers(tmp_path):
    doc = ds.fitz.open()
    for i, line in enumerate(("first page text", "second page ORIENTIR"), start=1):
        page = doc.new_page()
        page.insert_text((72, 72), line, fontsize=12)
    path = str(tmp_path / "p.pdf")
    doc.save(path)
    doc.close()

    res = ds.extract_text(path)
    assert res["unit"] == "page"
    assert res["status"] == "ok"

    db = str(tmp_path / "idx.db")
    ds.index_folder(str(tmp_path), False, db, lambda m: None, lambda r: None, _tr)
    hits = ds.search(db, "ORIENTIR", ds.MODE_PHRASE)[0]["hits"]
    assert hits[0]["label"] == 2               # მეორე გვერდზეა

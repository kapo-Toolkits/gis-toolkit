# -*- coding: utf-8 -*-
"""დოკუმენტებში ძებნის სუფთა ლოგიკა — GUI-სგან დამოუკიდებელი (ტესტირებადი).

აქ არ არის tkinter. სამი ნაწილი:

  1. **ამოღება** — ``extract_text()`` აბრუნებს დოკუმენტის სრულ ტექსტს და
     „ღუზების“ სიას (სიმბოლოს ოფსეტი → გვერდის/პარაგრაფის ნომერი), რომ
     ნაპოვნი ადგილი მერე დავალაგმოთ („გვ. 7“).
  2. **ინდექსი** — SQLite ქეში; ფაილი ხელახლა იკითხება მხოლოდ მაშინ, როცა
     mtime/ზომა შეიცვალა. პირველი ინდექსაცია ნელია, შემდეგი ძებნები მყისიერი.
  3. **ძებნა** — ფრაზა / ყველა სიტყვა (AND) / რომელიმე (OR), კონტექსტის
     ამონარიდით და მონიშვნის კოორდინატებით.

ქართულისთვის ძებნა **ქვესტრიქონულია** (სიტყვის საზღვარს არ ითხოვს): ენა
აგლუტინაციურია, ამიტომ „ნაკვეთ“ იპოვის „ნაკვეთი“-ს, „ნაკვეთის“-ს და
„ნაკვეთებში“-საც. სიტყვებს შორის ნებისმიერი რაოდენობის ჰარე/ხაზის გადატანა
დაიშვება, ანუ PDF-ში ორ ხაზად გატეხილი ფრაზაც იპოვება.

მძიმე პაკეტები (fitz/PyMuPDF, python-docx) „რბილად“ იმპორტდება — მოდული მათ
გარეშეც იმპორტადია (ტესტში/CI-ში), უბრალოდ შესაბამისი ფორმატი გამოტოვდება.
"""

import os
import re
import json
import time
import sqlite3
import traceback

try:                        # რბილი იმპორტი — PDF
    import pymupdf as fitz  # PyMuPDF >= 1.24
except ImportError:         # pragma: no cover
    try:
        import fitz         # ძველი PyMuPDF
    except ImportError:
        fitz = None

try:                        # რბილი იმპორტი — .docx
    import docx             # python-docx
except ImportError:         # pragma: no cover
    docx = None


# ---- პარამეტრები -----------------------------------------------------------
SUPPORTED_EXT = (".pdf", ".docx", ".txt", ".doc", ".rtf")

# როდის ჩავთვალოთ, რომ ფაილს ტექსტური ფენა არ აქვს (სავარაუდოდ დასკანერებულია).
# PDF-ს ვზომავთ **გვერდზე** გადაანგარიშებით: დასკანერებული გვერდი ~0 სიმბოლოს
# იძლევა, ნამდვილი ტექსტისა კი — ასეულებს. სხვა ფორმატებში მოკლე დოკუმენტი
# სავსებით ნორმალურია (ერთხაზიანი ცნობა), ამიტომ მხოლოდ სრულ სიცარიელეს ვიჭერთ.
MIN_CHARS_PER_PAGE = 8

CONTEXT = 150               # რამდენი სიმბოლო დამთხვევის ორივე მხარეს
MAX_SNIPPET_SPAN = 400      # ერთი ამონარიდი ამაზე განიერი არ გახდება
MAX_HITS_PER_FILE = 50      # UI-ის დასაცავად; საერთო რაოდენობა მაინც ითვლება

# ინდექსის სქემის ვერსია — ცვლილებისას ქეში თავიდან შენდება
SCHEMA_VERSION = 1

MODE_PHRASE = "phrase"
MODE_ALL = "all"
MODE_ANY = "any"


# ---- ტექსტის ამოღება -------------------------------------------------------
def _extract_pdf(path):
    """PDF → (ტექსტი, ღუზები, ერთეულის ტიპი). ღუზა = (ოფსეტი, გვერდის ნომერი)."""
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) is not installed")
    parts = []
    anchors = []
    offset = 0
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            t = page.get_text("text") or ""
            anchors.append((offset, i))
            parts.append(t)
            offset += len(t) + 1        # +1 — გამყოფი "\n"
    return "\n".join(parts), anchors, "page"


def _iter_docx_blocks(document):
    """დოკუმენტის სხეულის ელემენტები **დოკუმენტის რიგით** — პარაგრაფები და ცხრილები.

    ``document.paragraphs`` ცხრილების შიგთავსს ტოვებს, ``document.tables`` კი
    რიგს კარგავს; ამიტომ პირდაპირ XML სხეულს გავირბენთ.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def _docx_table_lines(table):
    """ცხრილის უჯრები → ტექსტის ხაზები (ჩადგმული ცხრილების ჩათვლით)."""
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            for nested in cell.tables:
                cell_text += " " + " ".join(_docx_table_lines(nested))
            cells.append(cell_text.strip())
        line = " | ".join(c for c in cells if c)
        if line.strip(" |"):
            lines.append(line)
    return lines


def _extract_docx(path):
    """.docx → (ტექსტი, ღუზები, ერთეულის ტიპი). ცხრილების უჯრებიც შედის.

    .docx-ს გვერდის ცნება არ აქვს (გვერდები Word-ის რენდერისას იბადება),
    ამიტომ ღუზა პარაგრაფის ნომერია.
    """
    if docx is None:
        raise RuntimeError("python-docx is not installed")
    document = docx.Document(path)
    parts = []
    anchors = []
    offset = 0
    n = 0
    for block in _iter_docx_blocks(document):
        if hasattr(block, "rows"):                     # ცხრილი
            lines = _docx_table_lines(block)
        else:                                          # პარაგრაფი
            lines = [block.text] if block.text.strip() else []
        for line in lines:
            n += 1
            anchors.append((offset, n))
            parts.append(line)
            offset += len(line) + 1
    return "\n".join(parts), anchors, "para"


def _extract_txt(path):
    """უბრალო ტექსტი — კოდირებას თანმიმდევრობით ვარჩევთ."""
    with open(path, "rb") as fh:
        raw = fh.read()
    for enc in ("utf-8-sig", "utf-8", "cp1251", "cp1252"):
        try:
            return raw.decode(enc), [(0, 1)], "para"
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace"), [(0, 1)], "para"


def _extract_word_com(path):
    """ძველი .doc / .rtf — Word-ის COM ავტომატიზაციით (მხოლოდ Windows + Word).

    Word ჩუმად იხსნება, ტექსტი მოაქვს და ისევ იხურება. Word რომ არ იყოს
    დაყენებული, ამომტყდება RuntimeError და ფაილი „გამოტოვებულად“ მოინიშნება.
    """
    try:
        import pythoncom
        import win32com.client
    except ImportError as e:                            # pragma: no cover
        raise RuntimeError("pywin32 is not available") from e

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True,
                                  AddToRecentFiles=False, Visible=False)
        text = doc.Content.Text or ""
    except Exception as e:                              # pragma: no cover
        raise RuntimeError(f"Word/COM: {e}") from e
    finally:
        try:
            if doc is not None:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()
    text = text.replace("\r", "\n")
    return text, [(0, 1)], "para"


EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
    ".doc": _extract_word_com,
    ".rtf": _extract_word_com,
}


def extract_text(path):
    """ფაილიდან ტექსტის ამოღება.

    აბრუნებს dict-ს: ``text``, ``anchors``, ``unit`` (page|para),
    ``status`` (ok|empty|error|skipped), ``note``.
    ``empty`` ნიშნავს, რომ ფაილი წაიკითხა, მაგრამ ტექსტი (თითქმის) არ იყო —
    დიდი ალბათობით დასკანერებული PDF-ია და OCR სჭირდება.
    """
    ext = os.path.splitext(path)[1].lower()
    fn = EXTRACTORS.get(ext)
    if fn is None:
        return {"text": "", "anchors": [], "unit": "para",
                "status": "skipped", "note": "unsupported: " + ext}
    try:
        text, anchors, unit = fn(path)
    except Exception as e:
        return {"text": "", "anchors": [], "unit": "para",
                "status": "error", "note": str(e)}

    n_chars = len(text.strip())
    if unit == "page" and anchors:
        # PDF — გვერდზე გადაანგარიშებული სიმკვრივე ამხელს დასკანერებულს
        empty = n_chars < MIN_CHARS_PER_PAGE * len(anchors)
    else:
        empty = n_chars == 0
    return {"text": text, "anchors": anchors, "unit": unit,
            "status": "empty" if empty else "ok", "note": ""}


# ---- ინდექსი (SQLite ქეში) -------------------------------------------------
def open_index(db_path):
    """ინდექსის ბაზის გახსნა/შექმნა."""
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS files (
            path       TEXT PRIMARY KEY,
            name       TEXT,
            ext        TEXT,
            mtime      REAL,
            size       INTEGER,
            text       TEXT,
            anchors    TEXT,
            unit       TEXT,
            status     TEXT,
            note       TEXT,
            indexed_at REAL
        )""")
    conn.execute("CREATE TABLE IF NOT EXISTS meta (key TEXT PRIMARY KEY, value TEXT)")
    row = conn.execute("SELECT value FROM meta WHERE key='schema'").fetchone()
    if row is None:
        conn.execute("INSERT INTO meta (key, value) VALUES ('schema', ?)",
                     (str(SCHEMA_VERSION),))
    elif row[0] != str(SCHEMA_VERSION):          # სქემა შეიცვალა — ქეში თავიდან
        conn.execute("DELETE FROM files")
        conn.execute("UPDATE meta SET value=? WHERE key='schema'", (str(SCHEMA_VERSION),))
    conn.commit()
    return conn


def iter_documents(folder, recursive=True):
    """საქაღალდის მხარდაჭერილი დოკუმენტები (დალაგებული, დროებითების გარეშე)."""
    found = []
    for root, dirs, names in os.walk(folder):
        dirs[:] = sorted(d for d in dirs if not d.startswith((".", "~$")))
        for name in sorted(names):
            if name.startswith("~$"):            # Word-ის დროებითი ფაილი
                continue
            if name.lower().endswith(SUPPORTED_EXT):
                found.append(os.path.join(root, name))
        if not recursive:
            break
    return found


def index_folder(folder, recursive, db_path, log, done, tr,
                 progress=None, cancel=None, force=False):
    """საქაღალდის ინდექსაცია (იგივე კონვენცია, რაც ``search_core.run_search``-ს).

    log(msg)             -> სტატუსის ჩაწერა
    done(result_dict)    -> დასრულებისას გამოძახება
    tr(key, **fmt)       -> თარგმანი
    progress(i, n)       -> პროგრესი (არჩევითი)
    cancel()             -> True თუ მომხმარებელმა გააუქმა (არჩევითი)
    force                -> True: ქეში იგნორირდება, ყველაფერი თავიდან იკითხება
    """
    conn = None
    try:
        if not folder or not os.path.isdir(folder):
            done({"error": tr("warn_folder")})
            return

        paths = iter_documents(folder, recursive)
        total = len(paths)
        if total == 0:
            done({"error": tr("no_docs")})
            return

        log(tr("idx_start", n=total))
        conn = open_index(db_path)

        cached = {r[0]: (r[1], r[2]) for r in
                  conn.execute("SELECT path, mtime, size FROM files")}

        n_new = n_upd = n_cached = n_empty = n_err = 0
        empty_files = []
        error_files = []

        for i, path in enumerate(paths, start=1):
            if cancel and cancel():
                conn.commit()
                done({"cancelled": True})
                return
            try:
                stat = os.stat(path)
            except OSError as e:
                n_err += 1
                error_files.append((path, str(e)))
                continue

            prev = cached.get(path)
            fresh = (prev is not None
                     and abs(prev[0] - stat.st_mtime) < 1e-6
                     and prev[1] == stat.st_size)
            if fresh and not force:
                n_cached += 1
            else:
                res = extract_text(path)
                conn.execute(
                    "INSERT OR REPLACE INTO files "
                    "(path, name, ext, mtime, size, text, anchors, unit, "
                    " status, note, indexed_at) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                    (path, os.path.basename(path),
                     os.path.splitext(path)[1].lower(),
                     stat.st_mtime, stat.st_size, res["text"],
                     json.dumps(res["anchors"]), res["unit"],
                     res["status"], res["note"], time.time()))
                if prev is None:
                    n_new += 1
                else:
                    n_upd += 1
                if res["status"] == "empty":
                    n_empty += 1
                    empty_files.append(path)
                elif res["status"] in ("error", "skipped"):
                    n_err += 1
                    error_files.append((path, res["note"]))

            if progress:
                progress(i, total)
            if i % 20 == 0 or i == total:
                log(tr("idx_progress", i=i, n=total))

        # ბაზიდან ამოვშალოთ ის ფაილები, რომლებიც საქაღალდეში აღარაა
        alive = set(paths)
        root = os.path.normpath(folder)
        stale = [p for p in cached
                 if os.path.normpath(p).startswith(root) and p not in alive]
        for p in stale:
            conn.execute("DELETE FROM files WHERE path=?", (p,))

        conn.commit()
        log(tr("idx_done", new=n_new, upd=n_upd, cached=n_cached))
        done({
            "total": total, "new": n_new, "updated": n_upd, "cached": n_cached,
            "empty": n_empty, "errors": n_err, "removed": len(stale),
            "empty_files": empty_files, "error_files": error_files,
        })
    except Exception as e:
        done({"error": "{}\n\n{}".format(e, traceback.format_exc())})
    finally:
        if conn is not None:
            conn.close()


# ---- ძებნა -----------------------------------------------------------------
def parse_query(query, mode):
    r"""საძიებო ტექსტი → რეგულარული გამოსახულებების სია.

    ფრაზა  — ერთი regex, სიტყვებს შორის ``\s+`` (ხაზის გადატანაც დაიშვება).
    AND/OR — თითო სიტყვაზე თითო regex; გამყოფად ჰარეც და „|“-ც მუშაობს.
    """
    query = (query or "").strip()
    if not query:
        return []
    if mode == MODE_PHRASE:
        words = query.split()
        if not words:
            return []
        pattern = r"\s+".join(re.escape(w) for w in words)
        return [re.compile(pattern, re.IGNORECASE)]
    words = [w for w in re.split(r"[|\s]+", query) if w]
    return [re.compile(re.escape(w), re.IGNORECASE) for w in words]


def anchor_for(anchors, offset):
    """ოფსეტს შეესაბამება ბოლო ღუზა, რომელიც მასზე ადრეა (გვერდის/პარაგრ. ნომერი)."""
    if not anchors:
        return None
    lo, hi = 0, len(anchors) - 1
    if offset < anchors[0][0]:
        return anchors[0][1]
    while lo < hi:
        mid = (lo + hi + 1) // 2
        if anchors[mid][0] <= offset:
            lo = mid
        else:
            hi = mid - 1
    return anchors[lo][1]


def group_spans(spans, context=CONTEXT, max_span=MAX_SNIPPET_SPAN):
    """ახლომდებარე დამთხვევების გაერთიანება — ერთ ამონარიდზე ერთი ჯგუფი.

    ერთ პარაგრაფში ორჯერ ნახსენები ფრაზა ორ **იდენტურ** ამონარიდს რომ არ
    დაბადოს, დამთხვევებს, რომლებიც ისედაც ერთმანეთის კონტექსტში ხვდება,
    ერთად ვაერთიანებთ და ბოლოს ყველას ვნიშნავთ.

    ჯგუფი ``max_span``-ზე განიერი არ ხდება — თორემ ტექსტში, სადაც სიტყვა
    ასჯერ ზედიზედ მეორდება, ერთი გაუთავებელი ამონარიდი დაიბადებოდა.
    """
    groups = []
    current = []
    for span in spans:
        if current and (span[0] - current[-1][1] > context
                        or span[1] - current[0][0] > max_span):
            groups.append(current)
            current = []
        current.append(span)
    if current:
        groups.append(current)
    return groups


def _collapse_ws(text, lo, hi):
    """[lo:hi] მონაკვეთში ჰარეების შეკუმშვა + რუკა ორიგინალის ინდექსებისთვის.

    აბრუნებს (შეკუმშული_ტექსტი, {ორიგინალის_ინდექსი: შეკუმშულის_ინდექსი}).
    რუკა საჭიროა, რომ მონიშვნა ზუსტად თავის დამთხვევაზე დადგეს — ტექსტით
    ძებნა აქ არ გამოდგება, რადგან ერთი და იგივე სიტყვა ამონარიდში
    რამდენჯერმე გვხვდება.
    """
    out = []
    pos = {}
    prev_space = False
    for i in range(lo, hi):
        ch = text[i]
        if ch.isspace():
            if prev_space:
                pos[i] = len(out) - 1
            else:
                pos[i] = len(out)
                out.append(" ")
                prev_space = True
        else:
            pos[i] = len(out)
            out.append(ch)
            prev_space = False
    return "".join(out), pos


def make_snippet(text, spans, context=CONTEXT):
    """ერთი ან რამდენიმე ახლომდებარე დამთხვევის ამონარიდი + მონიშვნები.

    ``spans`` — [(start, end), …] ერთი ჯგუფი (იხ. ``group_spans``).
    აბრუნებს (ამონარიდი, [(s, e), …]) — მონიშვნის კოორდინატები **ამონარიდის
    შიგნით**. კიდეები სიტყვის საზღვარზე სწორდება, რომ შუა ასოზე არ გაიჭრას.
    """
    if not spans:
        return "", []
    first, last = spans[0], spans[-1]
    lo = max(0, first[0] - context)
    hi = min(len(text), last[1] + context)
    if lo > 0:
        sp = text.find(" ", lo, first[0])
        if sp != -1:
            lo = sp + 1
    if hi < len(text):
        sp = text.rfind(" ", last[1], hi)
        if sp != -1:
            hi = sp

    collapsed, pos = _collapse_ws(text, lo, hi)
    shift = len(collapsed) - len(collapsed.lstrip())
    snippet = collapsed.strip()

    prefix = "…" if lo > 0 else ""
    suffix = "…" if hi < len(text) else ""
    pad = len(prefix) - shift

    hl = []
    for start, end in spans:
        if start not in pos or (end - 1) not in pos:
            continue
        s = pos[start] + pad
        e = pos[end - 1] + 1 + pad
        s = max(len(prefix), s)
        e = min(len(prefix) + len(snippet), e)
        if e > s:
            hl.append((s, e))
    return prefix + snippet + suffix, hl


def search(db_path, query, mode=MODE_PHRASE, max_hits_per_file=MAX_HITS_PER_FILE,
           cancel=None):
    """ინდექსში ძებნა → ფაილების სია დამთხვევებით.

    თითო ჩანაწერი:
        {path, name, ext, unit, status, count, truncated, hits:[{label, snippet, hl}]}
    ``count`` არის დამთხვევების **სრული** რაოდენობა, ``hits`` კი შეზღუდულია
    ``max_hits_per_file``-ით (``truncated`` აჩვენებს, მოიჭრა თუ არა).
    """
    regexes = parse_query(query, mode)
    if not regexes:
        return []

    conn = open_index(db_path)
    try:
        rows = conn.execute(
            "SELECT path, name, ext, text, anchors, unit, status FROM files "
            "WHERE status='ok' ORDER BY name COLLATE NOCASE").fetchall()
    finally:
        conn.close()

    results = []
    for path, name, ext, text, anchors_json, unit, status in rows:
        if cancel and cancel():
            break
        if not text:
            continue

        # AND — ყველა სიტყვა უნდა იყოს დოკუმენტში; სხვაგვარად ფაილს ვტოვებთ
        if mode == MODE_ALL and not all(rx.search(text) for rx in regexes):
            continue

        spans = []
        for rx in regexes:
            for m in rx.finditer(text):
                spans.append((m.start(), m.end()))
        if not spans:
            continue
        spans.sort()

        anchors = json.loads(anchors_json or "[]")
        groups = group_spans(spans)
        hits = []
        for group in groups[:max_hits_per_file]:
            snippet, hl = make_snippet(text, group)
            hits.append({"label": anchor_for(anchors, group[0][0]),
                         "snippet": snippet, "hl": hl, "start": group[0][0]})
        results.append({
            "path": path, "name": name, "ext": ext, "unit": unit,
            "status": status, "count": len(spans), "hits": hits,
            "truncated": len(groups) > max_hits_per_file,
        })

    results.sort(key=lambda r: (-r["count"], r["name"].lower()))
    return results


def index_stats(db_path):
    """ინდექსის მოკლე სტატისტიკა: სულ / ტექსტიანი / ცარიელი / შეცდომიანი."""
    if not os.path.exists(db_path):
        return {"total": 0, "ok": 0, "empty": 0, "error": 0, "empty_files": []}
    conn = open_index(db_path)
    try:
        counts = dict(conn.execute(
            "SELECT status, COUNT(*) FROM files GROUP BY status").fetchall())
        empty_files = [r[0] for r in conn.execute(
            "SELECT path FROM files WHERE status='empty' ORDER BY name").fetchall()]
    finally:
        conn.close()
    return {
        "total": sum(counts.values()),
        "ok": counts.get("ok", 0),
        "empty": counts.get("empty", 0),
        "error": counts.get("error", 0) + counts.get("skipped", 0),
        "empty_files": empty_files,
    }
